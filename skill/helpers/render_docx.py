"""Render clean resume markdown to an ATS-safe .docx.

ATS-safe template (plan / decisions): single column, one standard font, simple
bold headings, bullet lists — NO tables, text boxes, headers/footers, or columns,
which resume parsers mangle. Only call this on a CLEAN draft (tags.py reports
clean); tags are not stripped here.

Markdown supported: # name, ## section, ### subsection, - bullets, **bold**,
*italic*, and plain paragraphs. `---` and blank lines are skipped.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

FONT = "Calibri"
BODY_PT = 11


def _add_runs(paragraph, text: str):
    for part in re.split(r"(\*\*.+?\*\*|\*.+?\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def render(md_path, out_path):
    doc = Document()
    normal = doc.styles["Normal"].font
    normal.name = FONT
    normal.size = Pt(BODY_PT)

    for raw in Path(md_path).read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if not line.strip() or line.strip() == "---":
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(20)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(line[3:].strip())
            r.bold = True
            r.font.size = Pt(13)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(line[4:].strip())
            r.bold = True
            r.font.size = Pt(BODY_PT)
        elif re.match(r"^[-*] ", line):
            _add_runs(doc.add_paragraph(style="List Bullet"), line[2:].strip())
        else:
            _add_runs(doc.add_paragraph(), line.strip())

    out = Path(out_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


if __name__ == "__main__":
    import sys

    print(render(sys.argv[1], sys.argv[2]))
