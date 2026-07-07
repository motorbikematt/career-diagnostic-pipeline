import render_docx
from docx import Document


def test_render_ats_safe(tmp_path):
    md = tmp_path / "resume.md"
    md.write_text(
        "# Jane Candidate\n\n"
        "## Experience\n\n"
        "### Acme — Product Manager\n\n"
        "- Owned the **growth** roadmap and shipped features.\n"
        "- Ran *A/B testing* across the funnel.\n",
        encoding="utf-8",
    )
    out = tmp_path / "resume.docx"
    render_docx.render(md, out)
    assert out.exists()

    doc = Document(str(out))
    texts = [p.text for p in doc.paragraphs]
    assert "Jane Candidate" in texts
    assert any("growth" in t for t in texts)
    # ATS-safe: no tables in the rendered document.
    assert doc.tables == []
    # Bold inline run survived.
    growth_para = next(p for p in doc.paragraphs if "growth" in p.text)
    assert any(r.bold and "growth" in r.text for r in growth_para.runs)
